import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# Sample data (replace with your own)
data = {
    'Name': ['Alice', 'Bob', 'Charlie'],
    'Age': [24, 30, 22],
    'Department': ['Physics', 'Math', 'Biology']
}
df = pd.DataFrame(data)

# Create a figure and add a table
fig, ax = plt.subplots(figsize=(6, 2))
ax.axis('off')
table = ax.table(cellText=df.values,
                 colLabels=df.columns,
                 cellLoc='center',
                 loc='center')
table.auto_set_font_size(False)
table.set_fontsize(10)
table.scale(1.2, 1.5)  # scale width, height

# Save as PNG
png_filename = 'table_figure.png'
plt.savefig(png_filename, bbox_inches='tight')
plt.close(fig)

# Save to Word document
doc = Document()
doc.add_heading('Table Output', level=1)
doc.add_picture(png_filename, width=Inches(5.5))  # adjust width as needed
doc.save('table_output.docx')

print("Saved table as PNG and Word document.")